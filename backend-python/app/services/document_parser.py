import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> str:
    """解析 .docx 文件，提取段落和表格文本"""
    from docx import Document

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def parse_xlsx(file_path: str) -> str:
    """解析 .xlsx 文件，按行提取文本（只读模式，适合大文件）"""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue

        headers = [str(h).strip() if h else "" for h in rows[0]]
        parts.append(f"【工作表: {sheet}】")

        for row in rows[1:]:
            values = [str(v).strip() if v else "" for v in row]
            if not any(values):
                continue
            line_parts = []
            for h, v in zip(headers, values):
                if h and v:
                    line_parts.append(f"{h}: {v}")
            if line_parts:
                parts.append(", ".join(line_parts))

    wb.close()
    return "\n".join(parts)


def parse_file(file_path: str, file_type: str) -> str:
    """根据文件类型分发解析"""
    file_type = file_type.lower().strip(".")

    if file_type == "docx":
        return parse_docx(file_path)
    elif file_type == "xlsx":
        return parse_xlsx(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}，仅支持 docx/xlsx")
